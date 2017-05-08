import re

from django.contrib.auth.decorators import login_required
from django.forms import ModelForm
from django.http import JsonResponse
# Create your views here.
# this login required decorator is to not allow to any
# view without authenticating
from django.shortcuts import redirect
from django.shortcuts import render, get_object_or_404
from django.shortcuts import render_to_response
from django.template import RequestContext
from django.template.loader import render_to_string
from docx import Document

from reporting.forms import BookForm
from reporting.models import Company, AreaPersonalType


@login_required(login_url="login/")
def home(request):
    return render(request, "index.html")


@login_required(login_url="login/")
def form_view(request):
    return render(request, "form.html")


@login_required(login_url="login/")
def index_view(request):
    return render(request, "index.html")


@login_required(login_url="login/")
def dashboard_view(request):
    return render(request, "dashboard.html")


class AreaPersonalTypeForm(ModelForm):
    class Meta:
        model = AreaPersonalType
        fields = ['company', 'point', 'work_unit', 'method', 'flow_rate', 'media', 'technique',
                  'area_personal_type', 'classification_type', 'classification_value']


class CompanyForm(ModelForm):
    class Meta:
        model = Company
        fields = ['company_name', 'company_address',
                  'company_country', 'company_contact_no'
            , 'company_pic', 'consultant']


def company_list(request, template_name='company/company_list.html'):
    companys = Company.objects.all()
    data = {}
    data['object_list'] = companys
    return render(request, template_name, data)


# def company_create(request, template_name='company/company_form.html'):
#     form = CompanyForm(request.POST)
#     form2 = AreaPersonalTypeForm(request.POST)
#     if all([form.is_valid(), form2.is_valid()]):
#         form.save()
#         form2.save()
#         return redirect('company_list')
#     else:
#         form = CompanyForm(request.POST)
#         form2 = AreaPersonalTypeForm(request.POST)
#     return render(request, template_name, {'form': form, 'form2': form2})

def company_create(request, template_name='company/company_form.html'):
    form = CompanyForm(request.POST)

    if form.is_valid():
        form.save()
        return redirect('company_list')
    else:
        form = CompanyForm(request.POST or None)
    return render(request, template_name, {'form': form})


def company_update(request, pk, template_name='company/company_form.html'):
    company = get_object_or_404(Company, pk=pk)
    # cem = get_object_or_404(AreaPersonalType, pk=id)
    form = CompanyForm(request.POST or None, instance=company)
    cem_form = AreaPersonalTypeForm(request.POST)

    if all([form.is_valid(), cem_form.is_valid()]):
        form.save()
        cem_form.save()
        return redirect('company_list')
    return render(request, template_name, {'form': form, 'cem_form': cem_form})


def company_delete(request, pk, template_name='company/company_confirm_delete.html'):
    company = get_object_or_404(Company, pk=pk)
    if request.method == 'POST':
        company.delete()
        return redirect('company_list')
    return render(request, template_name, {'object': company})


# new code ############################################################################################


def cem_list(request, template_name='company/cem_list.html'):
    cems = AreaPersonalType.objects.all()
    data = {}
    data['object_list'] = cems
    return render(request, template_name, data)


def cem_create(request, template_name='company/cem_form.html'):
    form = AreaPersonalTypeForm(request.POST)

    if form.is_valid():
        form.save()
        return redirect('cem_list')
    else:
        form = AreaPersonalTypeForm(request.POST or None)
    return render(request, template_name, {'form': form})


def cem_update(request, pk, template_name='company/cem_form.html'):
    cem = get_object_or_404(AreaPersonalType, pk=pk)
    form = AreaPersonalTypeForm(request.POST or None, instance=cem)

    if form.is_valid():
        form.save()
        return redirect('cem_list')
    return render(request, template_name, {'form': form})


def cem_delete(request, pk, template_name='company/cem_confirm_delete.html'):
    cem = get_object_or_404(AreaPersonalType, pk=pk)
    if request.method == 'POST':
        cem.delete()
        return redirect('cem_list')
    return render(request, template_name, {'object': cem})


def docx_replace_regex(doc_obj, regex, replace):
    regex = re.compile(regex)

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


def initiate_replace(self, regex, replace):
    filename = "ENVOSHA/templates/CEM_Report.docx"
    doc = Document(filename)
    docx_replace_regex(doc, regex, replace)
    doc.save('ENVOSHA/static/CEM_Report_new.docx')

def cover_report(self,id):
    filename = "ENVOSHA/templates/CEM_Cover.docx"
    doc = Document(filename)
    comp = Company.objects.filter(id=id)
    for coms in comp:
        docx_replace_regex(doc, r"<company_name>", coms.company_name)
        docx_replace_regex(doc, r"<company_address>", coms.company_address)
    doc.save('ENVOSHA/static/CEM_Cover_new.docx')
    return render( self, 'company/cem_reports.html')


# regex1 = re.compile(r"Baseline")
# replace1 = r"{{ company_name }}"
# regex2 = re.compile(r"{{ naaaaaam }}")
# replace2 = r"{{ xna }} "
# filename = "ENVOSHA/templates/CEM_Report.docx"
# doc = Document(filename)
#
# docx_replace_regex(doc, regex1, replace1)
# docx_replace_regex(doc, regex1, replace1)
# docx_replace_regex(doc, regex2, replace2)
# doc.save('ENVOSHA/static/CEM_Report_new.docx')



def handler404(request):
    response = render_to_response('404.html', {},
                                  context_instance=RequestContext(request))
    response.status_code = 404
    return response


def handler500(request):
    response = render_to_response('500.html', {},
                                  context_instance=RequestContext(request))
    response.status_code = 500
    return response


def custom_404(request):
    return render(request, '404.html', {}, status=404)


def cem_reports(request, template_name='company/cem_reports.html'):
    companys = Company.objects.all()
    data = {}
    data['object_list'] = companys
    return render(request, template_name, data)


def book_list(request, pk):
    books = AreaPersonalType.objects.filter(company=pk, area_personal_type='area')
    area_mel = AreaPersonalType.objects.filter(company=pk, area_personal_type='area', classification_type='mel')
    area_cl = AreaPersonalType.objects.filter(company=pk, area_personal_type='area', classification_type='cl')
    area_twa = AreaPersonalType.objects.filter(company=pk, area_personal_type='area', classification_type='twa')
    personal = AreaPersonalType.objects.filter(company=pk, area_personal_type='personal')
    personal_mel = AreaPersonalType.objects.filter(company=pk, area_personal_type='personal', classification_type='mel')
    personal_cl = AreaPersonalType.objects.filter(company=pk, area_personal_type='personal', classification_type='cl')
    personal_twa = AreaPersonalType.objects.filter(company=pk, area_personal_type='personal', classification_type='twa')
    return render(request, 'books/book_list.html', {'books': books, 'personals': personal, 'area_mel': area_mel,
                                                    'area_cl': area_cl, 'area_twa': area_twa,
                                                    'personal_mel': personal_mel, 'personal_cl': personal_cl,
                                                    'personal_twa': personal_twa
                                                    })


def save_book_form(request, form, template_name):
    data = dict()
    if request.method == 'POST':
        if form.is_valid():
            form.save()
            data['form_is_valid'] = True
            books = AreaPersonalType.objects.all()
            data['html_book_list'] = render_to_string('books/includes/partial_book_list.html', {
                'books': books
            })
        else:
            data['form_is_valid'] = False
    context = {'form': form}
    data['html_form'] = render_to_string(template_name, context, request=request)
    return JsonResponse(data)


def book_create(request):
    if request.method == 'POST':
        form = BookForm(request.POST)
    else:
        form = BookForm()
    return save_book_form(request, form, 'books/includes/partial_book_create.html')


def book_update(request, pk):
    book = get_object_or_404(AreaPersonalType, pk=pk)
    if request.method == 'POST':
        form = BookForm(request.POST, instance=book)
    else:
        form = BookForm(instance=book)
    return save_book_form(request, form, 'books/includes/partial_book_update.html')


def book_delete(request, pk):
    book = get_object_or_404(AreaPersonalType, pk=pk)
    data = dict()
    if request.method == 'POST':
        book.delete()
        data['form_is_valid'] = True
        books = AreaPersonalType.objects.all()
        data['html_book_list'] = render_to_string('books/includes/partial_book_list.html', {
            'books': books
        })
    else:
        context = {'book': book}
        data['html_form'] = render_to_string('books/includes/partial_book_delete.html', context, request=request)
    return JsonResponse(data)


def comp_list(request, pk):
    comp = Company.objects.filter(id=pk)
    print comp

    return render(request, 'books/book_list.html', {'comp': comp})
