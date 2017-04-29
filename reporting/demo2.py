# coding=utf-8
import re

from docx import Document

from reporting.models import Company


def docx_replace_regex(doc_obj, regex, replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace)


company_name = Company.objects.filter(id=37)
print company_name

regex1 = re.compile(r"Baseline")
replace1 = r"motamail"
regex2 = re.compile(r"{{ naaaaaam }}")
replace2 = r"{{ xna }} "
filename = "ENVOSHA/templates/CEM_Report.docx"
doc = Document(filename)
docx_replace_regex(doc, regex1, replace1)
docx_replace_regex(doc, regex2, replace2)
doc.save('CEM_Report_new.docx')
